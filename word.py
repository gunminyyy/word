import streamlit as st
import pdfplumber
import re
from docx import Document
from datetime import datetime
import io

# --- 1. 텍스트 치환을 위한 헬퍼 함수 ---
def replace_text_in_doc(doc, replacements):
    """워드 문서 내의 단락과 표에서 지정된 텍스트를 찾아 바꿉니다."""
    # 일반 단락 검사
    for p in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)
    
    # 표(Table) 내부 검사
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in p.text:
                            p.text = p.text.replace(old_text, new_text)

# --- 2. Streamlit UI 구성 ---
st.title("📄 PDF to Word 자동 변환기")

# 상단: 좌우 2단 분할
col1, col2 = st.columns(2)

with col1:
    st.subheader("1. 원본 파일 업로드")
    uploaded_pdf = st.file_uploader("PDF 파일을 여기에 끌어다 놓으세요.", type=["pdf"])

with col2:
    st.subheader("2. 정보 입력 및 옵션")
    product_name = st.text_input("제품명")
    mode = st.selectbox("모드 선택", ["CFF", "HP", "HPD"])

st.divider()

# 하단: 좌우 2단 분할 (버튼 영역)
col3, col4 = st.columns(2)

with col3:
    convert_btn = st.button("변환 실행", use_container_width=True)

# --- 3. 데이터 추출 및 변환 로직 ---
if convert_btn:
    if not uploaded_pdf:
        st.error("원본 PDF 파일을 업로드해주세요.")
    elif not product_name:
        st.error("제품명을 입력해주세요.")
    else:
        with st.spinner("파일을 변환하는 중입니다..."):
            try:
                # 1. PDF 텍스트 추출
                pdf_text = ""
                with pdfplumber.open(uploaded_pdf) as pdf:
                    for page in pdf.pages:
                        extracted = page.extract_text()
                        if extracted:
                            pdf_text += extracted + "\n"
                
                replacements = {}
                
                # 2. 모드별 로직 (CFF 모드)
                if mode == "CFF":
                    # 제품명 치환
                    replacements["ESTHETIC AROMA B"] = product_name
                    
                    # COLOR 추출
                    color_match = re.search(r'COLOR\s*:(.*?)APPEARANCE\s*:', pdf_text, re.DOTALL | re.IGNORECASE)
                    if color_match:
                        color_val = color_match.group(1).strip().upper()
                        replacements["PALE YELLOW TO YELLOW"] = color_val
                    
                    # SPECIFIC GRAVITY 앞숫자 추출 및 계산
                    sg_match = re.search(r'SPECIFIC GRAVITY.*?\(\d+°C\)\s*:\s*([\d\.]+)\s*[±\+/-]\s*[\d\.]+', pdf_text, re.IGNORECASE)
                    if sg_match:
                        sg_base = float(sg_match.group(1))
                        sg_new_val = f"{sg_base - 0.01:.3f} ~ {sg_base + 0.01:.3f}"
                        replacements["0.902 ~ 0.922"] = sg_new_val
                        
                    # REFRACTIVE INDEX 앞숫자 추출 및 계산
                    ri_match = re.search(r'REFRACTIVE INDEX.*?\(\d+°C\)\s*:\s*([\d\.]+)\s*[±\+/-]\s*[\d\.]+', pdf_text, re.IGNORECASE)
                    if ri_match:
                        ri_base = float(ri_match.group(1))
                        ri_new_val = f"{ri_base - 0.01:.3f} ~ {ri_base + 0.01:.3f}"
                        replacements["1.466 ~ 1.476"] = ri_new_val
                        
                    # 날짜 변환 (예: 26. FEB. 2026)
                    current_date = datetime.now().strftime("%d. %b. %Y").upper()
                    replacements["07. OCT. 2024"] = current_date

                # 3. 워드 템플릿 불러오기 및 텍스트 치환
                # 깃허브에 올린 템플릿 파일 경로를 지정합니다.
                doc_path = "templates/company_form.docx"
                doc = Document(doc_path)
                
                replace_text_in_doc(doc, replacements)
                
                # 4. 결과물을 메모리 버퍼에 저장 (다운로드를 위해)
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                
                st.success("변환이 완료되었습니다! 우측에서 다운로드하세요.")
                
                # 우측 하단에 다운로드 버튼 표시
                with col4:
                    st.download_button(
                        label="결과물 다운로드 (.docx)",
                        data=bio,
                        file_name=f"{product_name}_변환결과.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            except Exception as e:
                st.error(f"오류가 발생했습니다: {e}")
                st.info("PDF 파일의 텍스트 구조가 예상과 다르거나, 양식 파일을 찾을 수 없는 경우일 수 있습니다.")
